"""Document processing service."""

import csv
import logging
import os
from datetime import datetime
from io import StringIO
from typing import List
from uuid import uuid4

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from config import get_settings
from rag.chunker import TextChunker
from rag.embeddings import EmbeddingService
from rag.vector_store import VectorStore

logger = logging.getLogger(__name__)


class DocumentService:
    """Handles document upload, parsing, chunking, and indexing."""

    def __init__(self) -> None:
        self.settings = get_settings()
        self.chunker = TextChunker(chunk_size=512, chunk_overlap=50)
        self.embedding_service = EmbeddingService()
        self.vector_store = VectorStore()

        os.makedirs(self.settings.upload_dir, exist_ok=True)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text content from a PDF file."""
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text content from a text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from a Word document."""
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)

    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text content from an HTML file."""
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)

    def extract_text_from_csv(self, file_path: str) -> str:
        """Extract text content from a CSV file."""
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = []
            for row in reader:
                rows.append(" | ".join(row))
            return "\n".join(rows)

    def extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from a file based on its extension."""
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext in [".txt", ".md", ".markdown"]:
            return self.extract_text_from_txt(file_path)
        elif ext == ".docx":
            return self.extract_text_from_docx(file_path)
        elif ext in [".html", ".htm"]:
            return self.extract_text_from_html(file_path)
        elif ext == ".csv":
            return self.extract_text_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def process_document(
        self,
        file_path: str,
        filename: str,
    ) -> dict:
        """
        Process a document: extract text, chunk, embed, and store.
        
        Returns summary of processing.
        """
        text = self.extract_text(file_path, filename)

        if not text.strip():
            raise ValueError("Could not extract any text from the document")

        chunks = self.chunker.split_text(text)
        logger.info(f"Created {len(chunks)} chunks from {filename}")

        embeddings = await self.embedding_service.embed_batch(chunks)
        logger.info(f"Generated {len(embeddings)} embeddings")

        chunk_ids = [str(uuid4()) for _ in chunks]
        metadatas = [
            {
                "filename": filename,
                "chunk_index": i,
                "uploaded_at": datetime.utcnow().isoformat(),
            }
            for i in range(len(chunks))
        ]

        self.vector_store.add_chunks(
            chunk_ids=chunk_ids,
            embeddings=embeddings,
            documents=chunks,
            metadatas=metadatas,
        )

        return {
            "filename": filename,
            "chunks_created": len(chunks),
            "text_length": len(text),
        }

    def get_documents(self) -> List[dict]:
        """Get list of all indexed documents."""
        return self.vector_store.get_all_documents()

    def delete_document(self, filename: str) -> None:
        """Delete a document and its chunks from the vector store."""
        self.vector_store.delete_by_filename(filename)

        file_path = os.path.join(self.settings.upload_dir, filename)
        if os.path.exists(file_path):
            os.remove(file_path)
